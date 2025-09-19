"""
Document Parser for Healthcare Software Requirements

Handles parsing of various document formats including PDFs, Word documents,
and XML healthcare specifications while preserving hierarchy and semantics.
"""

import os
import logging
from typing import Dict, List, Optional, Union, Any
from pathlib import Path
import PyPDF2
from docx import Document
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import pandas as pd

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for healthcare software requirement documents."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.xml', '.html', '.txt']
        
    def parse_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Parse a healthcare document and extract structured content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
            
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
            
        logger.info(f"Parsing document: {file_path}")
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_word(file_path)
            elif file_extension == '.xml':
                return self._parse_xml(file_path)
            elif file_extension == '.html':
                return self._parse_html(file_path)
            elif file_extension == '.txt':
                return self._parse_txt(file_path)
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise
            
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF documents."""
        content = []
        metadata = {}
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract metadata
            if pdf_reader.metadata:
                metadata = {
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'creator': pdf_reader.metadata.get('/Creator', ''),
                    'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                }
            
            # Extract text from all pages
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    content.append({
                        'page': page_num + 1,
                        'text': page_text,
                        'type': 'text'
                    })
                    
        return {
            'file_path': str(file_path),
            'file_type': 'pdf',
            'metadata': metadata,
            'content': content,
            'total_pages': len(content)
        }
        
    def _parse_word(self, file_path: Path) -> Dict[str, Any]:
        """Parse Word documents."""
        doc = Document(file_path)
        content = []
        
        # Extract paragraphs with hierarchy
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                style = paragraph.style.name if paragraph.style else 'Normal'
                content.append({
                    'index': i,
                    'text': paragraph.text,
                    'style': style,
                    'type': 'paragraph'
                })
                
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
                
            content.append({
                'index': table_idx,
                'data': table_data,
                'type': 'table'
            })
            
        return {
            'file_path': str(file_path),
            'file_type': 'word',
            'content': content,
            'total_paragraphs': len([c for c in content if c['type'] == 'paragraph']),
            'total_tables': len([c for c in content if c['type'] == 'table'])
        }
        
    def _parse_xml(self, file_path: Path) -> Dict[str, Any]:
        """Parse XML documents."""
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        content = self._extract_xml_content(root)
        
        return {
            'file_path': str(file_path),
            'file_type': 'xml',
            'content': content,
            'root_tag': root.tag
        }
        
    def _extract_xml_content(self, element: ET.Element, level: int = 0) -> List[Dict[str, Any]]:
        """Recursively extract content from XML elements."""
        content = []
        
        # Add current element
        if element.text and element.text.strip():
            content.append({
                'level': level,
                'tag': element.tag,
                'text': element.text.strip(),
                'attributes': dict(element.attrib),
                'type': 'text'
            })
            
        # Process children
        for child in element:
            content.extend(self._extract_xml_content(child, level + 1))
            
        return content
        
    def _parse_html(self, file_path: Path) -> Dict[str, Any]:
        """Parse HTML documents."""
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            
        content = []
        
        # Extract text from various HTML elements
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'span']):
            if element.get_text().strip():
                content.append({
                    'tag': element.name,
                    'text': element.get_text().strip(),
                    'type': 'html_element'
                })
                
        return {
            'file_path': str(file_path),
            'file_type': 'html',
            'content': content
        }
        
    def _parse_txt(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text documents."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
        # Split content into lines and create structured content
        lines = content.split('\n')
        structured_content = []
        
        for i, line in enumerate(lines):
            if line.strip():  # Only include non-empty lines
                structured_content.append({
                    'line_number': i + 1,
                    'text': line.strip(),
                    'type': 'text_line'
                })
                
        return {
            'file_path': str(file_path),
            'file_type': 'txt',
            'content': structured_content,
            'total_lines': len(structured_content)
        }
        
    def extract_clean_text(self, parsed_doc: Dict[str, Any]) -> str:
        """
        Extract clean, machine-readable text from parsed document.
        
        Args:
            parsed_doc: Parsed document dictionary
            
        Returns:
            Clean text content
        """
        text_parts = []
        
        for item in parsed_doc['content']:
            if item['type'] in ['text', 'paragraph', 'html_element', 'text_line']:
                text_parts.append(item['text'])
            elif item['type'] == 'table':
                # Convert table to readable format
                for row in item['data']:
                    text_parts.append(' | '.join(row))
                    
        return '\n'.join(text_parts)
        
    def extract_hierarchy(self, parsed_doc: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Extract hierarchical structure from parsed document.
        
        Args:
            parsed_doc: Parsed document dictionary
            
        Returns:
            List of hierarchical elements with levels
        """
        hierarchy = []
        
        for item in parsed_doc['content']:
            if item['type'] in ['paragraph', 'html_element']:
                level = 0
                
                # Determine hierarchy level based on style or tag
                if 'style' in item:
                    if 'heading' in item['style'].lower() or 'title' in item['style'].lower():
                        level = 1
                elif 'tag' in item and item['tag'] in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(item['tag'][1])
                    
                hierarchy.append({
                    'level': level,
                    'text': item['text'],
                    'type': item['type']
                })
                
        return hierarchy
